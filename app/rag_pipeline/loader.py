# app/rag_pipeline/loader.py
import tempfile
import aiohttp
from typing import List
from pathlib import Path

from langchain_text_splitters import SpacyTextSplitter
from app.config import settings

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

import spacy
import re

from docx import Document as DocxDocument

nlp = spacy.load("en_core_web_sm")
nlp.disable_pipes("lemmatizer")

async def download_to_tempfile(url: str, suffix: str) -> str:
    """Download Document from URL and save to a temp file"""
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    async with aiohttp.ClientSession() as sess:
        async with sess.get(url) as resp:
            async for chunk in resp.content.iter_chunked(8192):
                tmp.write(chunk)
    tmp.close()
    return tmp.name

def smart_split_insurance_document(text: str, chunk_size: int = 1500, chunk_overlap: int = 200) -> List[str]:

    section_pattern = (
        r"(\n\d+\.\d+\.?\s+[A-Z][^.\n]*(?:\s+means|\s+refers|\s+includes|\s+shall))"
    )
    sections = re.split(section_pattern, text)
    
    chunks = []
    current_chunk = ""
    
    for i, section in enumerate(sections):
        if not section.strip():
            continue
            
        # If adding this section would exceed chunk size, save current chunk
        if len(current_chunk) + len(section) > chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            # Start new chunk with overlap from previous
            if chunk_overlap > 0:
                overlap_text = current_chunk[-chunk_overlap:] if len(current_chunk) > chunk_overlap else current_chunk
                current_chunk = overlap_text + section
            else:
                current_chunk = section
        else:
            current_chunk += section
    
    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # If no sections found, fall back to regular splitting
    if len(chunks) <= 1:
        text_splitter = SpacyTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            pipeline = "en_core_web_sm",
            # separators=["\n\n", "\n", ". "]
        )
        text_splitter._tokenizer = nlp
        text_splitter._tokenizer.disable_pipes("lemmatizer")
        return text_splitter.split_text(text)
    
    # Post-process: ensure no chunk exceeds maximum safe size (6000 chars ≈ 1500 tokens)
    final_chunks = []
    max_safe_size = 6000
    
    for chunk in chunks:
        if len(chunk) <= max_safe_size:
            final_chunks.append(chunk)
        else:
            # Split oversized chunks using recursive splitter
            text_splitter = SpacyTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                pipeline = "en_core_web_sm",
                # separators=["\n\n", "\n", ". "]
            )
            
            text_splitter._tokenizer = nlp
            text_splitter._tokenizer.disable_pipes("lemmatizer")
            sub_chunks = text_splitter.split_text(chunk)
            final_chunks.extend(sub_chunks)
    
    return final_chunks

async def load_and_chunk_pdf(url: str) -> List[Document]:
    """Load PDF from URL and chunk it into documents"""
    # Download PDF to temp file
    temp_path = await download_to_tempfile(url, ".pdf")
    
    try:
        # Load PDF
        loader = PyPDFLoader(temp_path)
        raw_documents = loader.load()
        
        # Combine all pages into one text for better section detection
        full_text = "\n".join([doc.page_content for doc in raw_documents])
        
        # Use custom smart splitting
        chunk_texts = smart_split_insurance_document(
            full_text, 
            chunk_size=settings.CHUNK_SIZE, 
            chunk_overlap=settings.CHUNK_OVERLAP,
        )
        
        # Convert back to Document objects
        chunks = []
        for i, chunk_text in enumerate(chunk_texts):
            # Find which page this chunk primarily belongs to
            page_num = 0
            char_count = 0
            for j, raw_doc in enumerate(raw_documents):
                if char_count + len(raw_doc.page_content) >= len(chunk_text) * 0.5:
                    page_num = j
                    break
                char_count += len(raw_doc.page_content)
            
            doc = Document(
                page_content=chunk_text,
                metadata={
                    "chunk_index": i,
                    "source": temp_path,
                    "page": page_num
                }
            )
            chunks.append(doc)
        
        return chunks
        
    finally:
        # Clean up temp file
        import os
        if os.path.exists(temp_path):
            os.unlink(temp_path)

async def load_and_chunk_docx(url: str) -> List[Document]:
    """Load DOCX from URL and chunk it into documents."""
    temp_path = await download_to_tempfile(url, ".docx")
    try:

        doc = DocxDocument(temp_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        chunk_texts = smart_split_insurance_document(
            full_text,
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
        )

        return [
            Document(
                page_content=chunk_text,
                metadata={"chunk_index": i, "source": temp_path},
            )
            for i, chunk_text in enumerate(chunk_texts)
        ]
    finally:
        import os

        if os.path.exists(temp_path):
            os.unlink(temp_path)


async def load_and_chunk_image(url: str) -> List[Document]:
    """Load image from URL and perform OCR before chunking."""
    suffix = Path(url).suffix or ".png"
    temp_path = await download_to_tempfile(url, suffix)
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(temp_path)
        full_text = pytesseract.image_to_string(image)

        chunk_texts = smart_split_insurance_document(
            full_text,
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
        )

        return [
            Document(
                page_content=chunk_text,
                metadata={"chunk_index": i, "source": temp_path},
            )
            for i, chunk_text in enumerate(chunk_texts)
        ]
    finally:
        import os

        if os.path.exists(temp_path):
            os.unlink(temp_path)


async def load_and_chunk_eml(url: str) -> List[Document]:
    """Load EML email from URL and chunk its text content."""
    temp_path = await download_to_tempfile(url, ".eml")
    try:
        from email import policy
        from email.parser import BytesParser

        with open(temp_path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)

        if msg.is_multipart():
            parts = [
                part.get_content()
                for part in msg.walk()
                if part.get_content_type() == "text/plain"
            ]
            full_text = "\n".join(parts)
        else:
            full_text = msg.get_content()

        chunk_texts = smart_split_insurance_document(
            full_text,
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
        )

        return [
            Document(
                page_content=chunk_text,
                metadata={"chunk_index": i, "source": temp_path},
            )
            for i, chunk_text in enumerate(chunk_texts)
        ]
    finally:
        import os

        if os.path.exists(temp_path):
            os.unlink(temp_path)

from urllib.parse import urlparse
from pathlib import Path

def get_file_ext(url: str) -> str:
    parsed = urlparse(url)
    path = parsed.path  # extracts "/foo/bar/file.pdf"
    return Path(path).suffix.lower()


async def load_and_chunk(url: str) -> List[Document]:
    """Dispatch loading and chunking based on file extension."""
    ext = get_file_ext(url)
    if ext == ".pdf":
        return await load_and_chunk_pdf(url)
    if ext == ".docx":
        return await load_and_chunk_docx(url)
    if ext in {".png", ".jpg", ".jpeg"}:
        return await load_and_chunk_image(url)
    if ext == ".eml":
        return await load_and_chunk_eml(url)
    raise ValueError(f"Unsupported file type: {ext}")